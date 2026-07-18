from __future__ import annotations

import hashlib
import json
import mimetypes
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import unquote, urlsplit

from legal_assistant.domain.models import LegalDocument


class UnsupportedDocumentFormatError(ValueError):
    """Raised when a local document has no registered extraction strategy."""


class LocalFileDocumentParser:
    """Parse common local files into the application's document contract.

    Heavy format libraries are imported lazily, so JSONL remains usable without
    installing PDF, Word, or spreadsheet dependencies.
    """

    SUPPORTED_EXTENSIONS = frozenset(
        {".pdf", ".docx", ".xlsx", ".xlsm", ".xltx", ".xltm", ".xls", ".jsonl"}
    )

    def __init__(self, *, jurisdiction: str = "IR") -> None:
        self._jurisdiction = jurisdiction

    def parse(self, source_uri: str) -> list[LegalDocument]:
        path = self._resolve_path(source_uri)
        suffix = path.suffix.casefold()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            supported = ", ".join(sorted(self.SUPPORTED_EXTENSIONS))
            raise UnsupportedDocumentFormatError(
                f"Unsupported document format {suffix or '<none>'!r}: {path}. "
                f"Supported formats: {supported}"
            )
        if not path.is_file():
            raise FileNotFoundError(f"Document source does not exist: {path}")

        if suffix == ".pdf":
            return [self._parse_pdf(path)]
        if suffix == ".docx":
            return [self._parse_docx(path)]
        if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
            return self._parse_openpyxl(path)
        if suffix == ".xls":
            return self._parse_xls(path)
        return self._parse_jsonl(path)

    @staticmethod
    def _resolve_path(source_uri: str) -> Path:
        parts = urlsplit(source_uri)
        if parts.scheme and parts.scheme != "file":
            raise ValueError(f"LocalFileDocumentParser only accepts local paths: {source_uri}")
        raw_path = unquote(parts.path) if parts.scheme == "file" else source_uri
        return Path(raw_path).expanduser().resolve()

    @staticmethod
    def _document_id(path: Path, discriminator: str = "") -> str:
        identity = f"{path.resolve()}#{discriminator}"
        return f"document:{hashlib.sha256(identity.encode('utf-8')).hexdigest()[:24]}"

    @staticmethod
    def _source_metadata(path: Path, source_format: str) -> dict[str, Any]:
        mime_type, _ = mimetypes.guess_type(path.name)
        return {
            "source_format": source_format,
            "mime_type": mime_type or "application/octet-stream",
            "file_name": path.name,
            "file_size": path.stat().st_size,
        }

    def _build_document(
        self,
        path: Path,
        *,
        text: str,
        source_format: str,
        discriminator: str = "",
        title: str | None = None,
        source_uri: str | None = None,
        document_type: str = "document",
        metadata: dict[str, Any] | None = None,
        effective_date: str | None = None,
        publication_date: str | None = None,
        version: str | None = None,
        jurisdiction: str | None = None,
    ) -> LegalDocument:
        combined_metadata = self._source_metadata(path, source_format)
        combined_metadata.update(metadata or {})
        return LegalDocument(
            id=self._document_id(path, discriminator),
            title=title or path.stem,
            source_uri=source_uri or path.as_uri(),
            jurisdiction=jurisdiction or self._jurisdiction,
            document_type=document_type,
            text=text.strip(),
            effective_date=effective_date,
            publication_date=publication_date,
            version=version,
            parser_name="local_file",
            metadata=combined_metadata,
        )

    def _parse_pdf(self, path: Path) -> LegalDocument:
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - depends on optional extra
            raise RuntimeError("PDF support requires the 'ingestion' extra (pypdf).") from exc

        reader = PdfReader(str(path))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        return self._build_document(
            path,
            text="\n\n".join(page for page in pages if page),
            source_format="pdf",
            metadata={"page_count": len(reader.pages)},
        )

    def _parse_docx(self, path: Path) -> LegalDocument:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:  # pragma: no cover - depends on optional extra
            raise RuntimeError("Word support requires the 'ingestion' extra (python-docx).") from exc

        document = DocxDocument(str(path))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    blocks.append("\t".join(values))
        return self._build_document(
            path,
            text="\n".join(blocks),
            source_format="docx",
            metadata={"paragraph_count": len(document.paragraphs), "table_count": len(document.tables)},
        )

    def _parse_openpyxl(self, path: Path) -> list[LegalDocument]:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:  # pragma: no cover - depends on optional extra
            raise RuntimeError("Excel support requires the 'ingestion' extra (openpyxl).") from exc

        workbook = load_workbook(filename=path, read_only=True, data_only=True)
        try:
            return [
                self._spreadsheet_document(
                    path,
                    sheet.title,
                    ([cell.value for cell in row] for row in sheet.iter_rows()),
                    source_format=path.suffix.casefold().lstrip("."),
                )
                for sheet in workbook.worksheets
            ]
        finally:
            workbook.close()

    def _parse_xls(self, path: Path) -> list[LegalDocument]:
        try:
            import xlrd
        except ImportError as exc:  # pragma: no cover - depends on optional extra
            raise RuntimeError("Legacy Excel support requires the 'ingestion' extra (xlrd).") from exc

        workbook = xlrd.open_workbook(str(path), on_demand=True)
        try:
            documents: list[LegalDocument] = []
            for sheet in workbook.sheets():
                rows = (sheet.row_values(index) for index in range(sheet.nrows))
                documents.append(
                    self._spreadsheet_document(
                        path, sheet.name, rows, source_format="xls"
                    )
                )
            return documents
        finally:
            workbook.release_resources()

    def _spreadsheet_document(
        self,
        path: Path,
        sheet_name: str,
        rows: Iterable[Iterable[Any]],
        *,
        source_format: str,
    ) -> LegalDocument:
        lines: list[str] = []
        row_count = 0
        for row_count, row in enumerate(rows, start=1):
            values = ["" if value is None else str(value).strip() for value in row]
            while values and not values[-1]:
                values.pop()
            if any(values):
                lines.append("\t".join(values))
        return self._build_document(
            path,
            text="\n".join(lines),
            source_format=source_format,
            discriminator=sheet_name,
            title=f"{path.stem} — {sheet_name}",
            metadata={"sheet_name": sheet_name, "row_count": row_count},
        )

    def _parse_jsonl(self, path: Path) -> list[LegalDocument]:
        documents: list[LegalDocument] = []
        with path.open("r", encoding="utf-8-sig") as handle:
            for record_index, line in enumerate(handle, start=1):
                if not line.strip():
                    continue
                record = json.loads(line)
                if not isinstance(record, dict):
                    raise ValueError(f"JSONL record {record_index} must be an object: {path}")
                text = next(
                    (str(record[key]) for key in ("text", "content", "body") if record.get(key) is not None),
                    "",
                )
                if not text.strip():
                    raise ValueError(f"JSONL record {record_index} has no text/content/body: {path}")
                reserved = {
                    "id", "document_id", "title", "text", "content", "body", "url", "source_uri",
                    "jurisdiction", "document_type", "effective_date", "publication_date", "version", "metadata",
                }
                metadata = dict(record.get("metadata") or {})
                metadata.update({key: value for key, value in record.items() if key not in reserved})
                metadata["record_index"] = record_index
                discriminator = str(record.get("document_id") or record.get("id") or record_index)
                documents.append(
                    self._build_document(
                        path,
                        text=text,
                        source_format="jsonl",
                        discriminator=discriminator,
                        title=str(record.get("title") or f"{path.stem} #{record_index}"),
                        source_uri=str(record.get("source_uri") or record.get("url") or f"{path.as_uri()}#L{record_index}"),
                        document_type=str(record.get("document_type") or "document"),
                        metadata=metadata,
                        effective_date=self._optional_string(record.get("effective_date")),
                        publication_date=self._optional_string(record.get("publication_date")),
                        version=self._optional_string(record.get("version")),
                        jurisdiction=str(record.get("jurisdiction") or self._jurisdiction),
                    )
                )
        return documents

    @staticmethod
    def _optional_string(value: Any) -> str | None:
        return None if value is None else str(value)
